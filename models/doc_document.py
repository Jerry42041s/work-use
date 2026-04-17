import base64
import io
import re
import subprocess
import tempfile
import os
from odoo import models, fields, api
from odoo.exceptions import UserError


# ─── python-docx HTML 轉換輔助函式 ────────────────────────────────────────────

def _add_runs_from_node(para, node):
    """將節點的文字（含行內格式）加入 docx paragraph。"""
    if node.text:
        para.add_run(node.text)
    for child in node:
        ctag = (child.tag or '').lower() if isinstance(child.tag, str) else ''
        if ctag in ('strong', 'b'):
            run = para.add_run(child.text_content())
            run.bold = True
        elif ctag in ('em', 'i'):
            run = para.add_run(child.text_content())
            run.italic = True
        elif ctag == 'u':
            run = para.add_run(child.text_content())
            run.underline = True
        elif ctag == 'br':
            para.add_run('\n')
        elif ctag == 'span':
            style_str = child.get('style', '')
            run = para.add_run(child.text_content())
            if 'bold' in style_str:
                run.bold = True
            if 'italic' in style_str:
                run.italic = True
        else:
            text = child.text_content()
            if text:
                para.add_run(text)
        if child.tail:
            para.add_run(child.tail)


def _add_table_to_docx(doc, node):
    """將 HTML table 節點轉換為 python-docx 表格。"""
    rows = node.xpath('.//tr')
    if not rows:
        return
    max_cols = max(
        sum(int(td.get('colspan', 1)) for td in row.xpath('.//td|.//th'))
        for row in rows
    ) or 1
    table = doc.add_table(rows=len(rows), cols=max_cols)
    try:
        table.style = 'Table Grid'
    except Exception:
        pass
    for r_idx, row in enumerate(rows):
        cells = row.xpath('.//td|.//th')
        for c_idx, cell in enumerate(cells):
            if c_idx < max_cols:
                try:
                    table.cell(r_idx, c_idx).text = (cell.text_content() or '').strip()
                except Exception:
                    pass


def _html_node_to_docx(doc, node):
    """遞迴將 lxml HTML 節點轉換為 python-docx 結構。"""
    tag = (node.tag or '').lower() if isinstance(node.tag, str) else ''

    if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
        level = int(tag[1])
        para = doc.add_heading(level=level)
        _add_runs_from_node(para, node)

    elif tag == 'table':
        _add_table_to_docx(doc, node)

    elif tag == 'hr':
        doc.add_paragraph('─' * 40)

    elif tag in ('ul', 'ol'):
        for child in node:
            if (child.tag or '').lower() == 'li':
                para = doc.add_paragraph(style='List Bullet')
                _add_runs_from_node(para, child)

    elif tag in ('p', 'li', 'blockquote', 'pre'):
        text = (node.text_content() or '').strip()
        if text:
            para = doc.add_paragraph()
            _add_runs_from_node(para, node)

    elif tag in ('div', 'section', 'article', 'main', 'aside',
                 'header', 'footer', 'body', 'span'):
        # 若 div 有直接文字內容且無塊級子節點，作為段落
        block_tags = {'p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
                      'table', 'ul', 'ol', 'hr', 'div', 'blockquote', 'pre'}
        has_block_children = any(
            (c.tag or '').lower() in block_tags for c in node
        )
        if not has_block_children:
            text = (node.text_content() or '').strip()
            if text:
                para = doc.add_paragraph()
                _add_runs_from_node(para, node)
        else:
            for child in node:
                _html_node_to_docx(doc, child)

    # script/style/meta 等略過

# 超過此大小（bytes）自動存入 ir.attachment
CONTENT_SIZE_THRESHOLD = 500 * 1024  # 500 KB

import logging
_logger = logging.getLogger(__name__)


def _fix_docx_table_widths(docx_bytes, usable_w_mm):
    """修正 LibreOffice 生成的 DOCX 中所有表格寬度（含巢狀），使其符合頁面可用寬度。

    LO 24.2 以內建預設邊距（~5mm）計算表格寬，導致表格寬 200mm 而版心只有 170mm。
    CSS / inline style / !important 全部無效；唯一可靠解法是在 OOXML 層直接修正。

    修改範圍（三層同步）：
    1. <w:tblGrid> / <w:gridCol> — 欄寬網格（等比縮放）
    2. <w:tblW>                  — 表格總寬（設為 target_twips, type=dxa）
    3. <w:tcW>                   — 儲存格寬度（type=dxa 者等比縮放）

    失敗保護：任何異常回傳原始 bytes（優雅降級，確保使用者至少能下載檔案）。
    """
    try:
        from lxml import etree
        import zipfile

        WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

        def W(tag):
            return f'{{{WNS}}}{tag}'

        # mm → twips（1 inch = 1440 twips；1 inch = 25.4 mm）
        target_twips = round(usable_w_mm * 1440 / 25.4)

        # 只讀取 document.xml，其餘 DOCX 內容不解析（省記憶體）
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zin:
            doc_xml_bytes = zin.read('word/document.xml')

        root = etree.fromstring(doc_xml_bytes)

        # 找出所有 w:tbl（含巢狀），lxml 的 findall 以文件順序回傳（前序）
        all_tbls = root.findall(f'.//{W("tbl")}')

        for tbl in all_tbls:
            # ── Step 1: 讀取並等比縮放 <w:tblGrid> ──────────────────────────
            tbl_grid = tbl.find(W('tblGrid'))
            grid_cols = tbl_grid.findall(W('gridCol')) if tbl_grid is not None else []

            current_total = (
                sum(int(c.get(W('w'), 0)) for c in grid_cols)
                if grid_cols else 0
            )

            # current_total 為 0（空表格）或已在合理範圍（±2%）則跳過
            if current_total <= 0:
                continue
            if abs(current_total - target_twips) / target_twips < 0.02:
                continue

            scale = target_twips / current_total

            # 等比縮放 gridCol，最後一欄補足捨入誤差
            total_assigned = 0
            for i, col in enumerate(grid_cols):
                if i < len(grid_cols) - 1:
                    new_w = max(1, round(int(col.get(W('w'), 0)) * scale))
                else:
                    new_w = max(1, target_twips - total_assigned)
                col.set(W('w'), str(new_w))
                total_assigned += new_w

            # ── Step 2: 修正 <w:tblW> ────────────────────────────────────────
            tbl_pr = tbl.find(W('tblPr'))
            if tbl_pr is not None:
                tbl_w = tbl_pr.find(W('tblW'))
                if tbl_w is None:
                    tbl_w = etree.SubElement(tbl_pr, W('tblW'))
                tbl_w.set(W('w'), str(target_twips))
                tbl_w.set(W('type'), 'dxa')

            # ── Step 3: 等比縮放所有 <w:tcW>（type=dxa）──────────────────────
            for tc in tbl.iter(W('tc')):
                tc_pr = tc.find(W('tcPr'))
                if tc_pr is None:
                    continue
                tc_w = tc_pr.find(W('tcW'))
                if tc_w is None or tc_w.get(W('type')) != 'dxa':
                    continue
                old_w = int(tc_w.get(W('w'), 0))
                if old_w <= 0:
                    continue
                tc_w.set(W('w'), str(max(1, round(old_w * scale))))

        # ── 重新打包 DOCX（只替換 document.xml，其餘保持原樣）────────────────
        modified_xml = etree.tostring(
            root, xml_declaration=True, encoding='UTF-8', standalone=True
        )
        output = io.BytesIO()
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zin:
            with zipfile.ZipFile(output, 'w', zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = (
                        modified_xml
                        if item.filename == 'word/document.xml'
                        else zin.read(item.filename)
                    )
                    zout.writestr(item, data)
        output.seek(0)
        return output.read()

    except Exception as e:
        _logger.warning('[DocEditor] _fix_docx_table_widths 失敗，回傳原始 DOCX: %s', e)
        return docx_bytes  # 優雅降級：確保使用者至少可下載（邊距可能有誤但檔案可開啟）


class DocDocument(models.Model):
    _name = 'doc.document'
    _description = '文件'
    _inherit = ['mail.thread', 'mail.activity.mixin', 'doc.render.mixin']
    _order = 'write_date desc'

    name = fields.Char(string='文件名稱', required=True, default='未命名文件', tracking=True)

    # 內容欄位（sanitize=False，由 Server-Side Sanitizer 在 create/write 中處理）
    content_html = fields.Html(
        string='內容',
        sanitize=False,
        sanitize_attributes=False,
    )
    header_html = fields.Html(
        string='頁首',
        sanitize=False,
        sanitize_attributes=False,
    )
    footer_html = fields.Html(
        string='頁尾',
        sanitize=False,
        sanitize_attributes=False,
    )

    # 文件設定
    model_id = fields.Many2one(
        'ir.model',
        string='關聯模型',
        ondelete='set null',
        help='用於欄位變數渲染的目標模型',
    )
    template_id = fields.Many2one(
        'doc.template',
        string='使用範本',
        ondelete='set null',
    )
    page_format = fields.Selection([
        ('A4', 'A4'),
        ('A3', 'A3'),
        ('A5', 'A5'),
        ('letter', 'Letter'),
        ('legal', 'Legal'),
    ], string='頁面格式', default='A4')
    margin_top = fields.Integer(string='上邊距 (px)', default=96)
    margin_bottom = fields.Integer(string='下邊距 (px)', default=96)
    margin_left = fields.Integer(string='左邊距 (px)', default=96)
    margin_right = fields.Integer(string='右邊距 (px)', default=96)
    has_different_first_page = fields.Boolean(string='首頁不同頁首/頁尾', default=False)
    first_header_html = fields.Html(string='首頁頁首', sanitize=False)
    first_footer_html = fields.Html(string='首頁頁尾', sanitize=False)

    # 協作者
    collaborator_ids = fields.Many2many(
        'res.users',
        'doc_document_collaborator_rel',
        'document_id', 'user_id',
        string='協作者',
    )

    # 多公司隔離
    company_id = fields.Many2one(
        'res.company',
        string='公司',
        required=True,
        default=lambda self: self.env.company,
        index=True,
    )

    # 大文件 attachment 策略（>500KB 自動存入 ir.attachment）
    content_attachment_id = fields.Many2one(
        'ir.attachment',
        string='內容附件（大文件）',
        ondelete='set null',
    )
    is_large_document = fields.Boolean(string='大文件', default=False)

    # 多欄排版預設值
    default_column_count = fields.Selection([
        ('1', '1 欄'), ('2', '2 欄'), ('3', '3 欄'),
    ], string='預設欄數', default='1')
    default_column_gap = fields.Integer(string='欄間距 (px)', default=24)
    column_rule_style = fields.Selection([
        ('none', '無'), ('solid', '實線'), ('dashed', '虛線'), ('dotted', '點線'),
    ], string='欄分隔線', default='none')

    # 需要 sanitize 的 HTML 欄位清單
    _HTML_FIELDS = ('content_html', 'header_html', 'footer_html',
                    'first_header_html', 'first_footer_html')

    # ─── CRUD Hooks ─────────────────────────────────────────────────

    @api.model_create_multi
    def create(self, vals_list):
        sanitizer = self.env['doc.sanitizer']
        for vals in vals_list:
            for field in self._HTML_FIELDS:
                if vals.get(field):
                    vals[field] = sanitizer.sanitize_html(vals[field])
        return super().create(vals_list)

    def write(self, vals):
        sanitizer = self.env['doc.sanitizer']
        for field in self._HTML_FIELDS:
            if vals.get(field):
                vals[field] = sanitizer.sanitize_html(vals[field])

        # 大文件自動存入 attachment
        if 'content_html' in vals and vals['content_html']:
            html = vals['content_html']
            if len(html.encode('utf-8')) > CONTENT_SIZE_THRESHOLD:
                # 各筆記錄分別建立/更新 attachment，不共用 vals
                base_vals = dict(vals)
                base_vals.pop('content_html', None)
                base_vals['is_large_document'] = True
                base_vals['content_html'] = ''
                for rec in self:
                    att = rec._store_content_as_attachment(html)
                    super(DocDocument, rec).write(
                        dict(base_vals, content_attachment_id=att.id)
                    )
                return True
            else:
                # 文件縮小：清除舊 attachment
                vals['is_large_document'] = False
                vals['content_attachment_id'] = False

        return super().write(vals)

    def _store_content_as_attachment(self, html):
        """將大文件內容存入 ir.attachment。"""
        self.ensure_one()
        data = base64.b64encode(html.encode('utf-8'))
        if self.content_attachment_id:
            self.content_attachment_id.write({'datas': data})
            return self.content_attachment_id
        return self.env['ir.attachment'].create({
            'name': f'doc_content_{self.id}.html',
            'type': 'binary',
            'datas': data,
            'res_model': 'doc.document',
            'res_id': self.id,
        })

    def get_content_html(self):
        """統一取得 HTML 內容（自動判斷來源）。"""
        self.ensure_one()
        if self.is_large_document and self.content_attachment_id:
            return base64.b64decode(self.content_attachment_id.datas).decode('utf-8')
        return self.content_html or ''

    # ─── Actions ─────────────────────────────────────────────────────

    def action_open_editor(self):
        """開啟全螢幕文件編輯器。"""
        self.ensure_one()
        return {
            'type': 'ir.actions.client',
            'tag': 'dobtor_doc_editor.action_doc_editor',
            'context': {
                'doc_id': self.id,
                'doc_name': self.name,
            },
            'target': 'fullscreen',
        }

    def action_export_pdf(self):
        """匯出 PDF 並下載。"""
        self.ensure_one()
        pdf_bytes = self._generate_pdf()
        attachment = self.env['ir.attachment'].create({
            'name': f'{self.name}.pdf',
            'type': 'binary',
            'datas': base64.b64encode(pdf_bytes),
            'res_model': 'doc.document',
            'res_id': self.id,
            'mimetype': 'application/pdf',
        })
        return {
            'type': 'ir.actions.act_url',
            'url': f'/web/content/{attachment.id}?download=true',
            'target': 'new',
        }

    def action_export_docx(self):
        """匯出 DOCX 並下載。"""
        self.ensure_one()
        docx_bytes = self._generate_docx_via_libreoffice()
        attachment = self.env['ir.attachment'].create({
            'name': f'{self.name}.docx',
            'type': 'binary',
            'datas': base64.b64encode(docx_bytes),
            'res_model': 'doc.document',
            'res_id': self.id,
            'mimetype': ('application/vnd.openxmlformats-officedocument'
                         '.wordprocessingml.document'),
        })
        return {
            'type': 'ir.actions.act_url',
            'url': f'/web/content/{attachment.id}?download=true',
            'target': 'new',
        }

    # ─── 後端匯出邏輯 ─────────────────────────────────────────────────

    def _build_full_html(self, rendered_body=None):
        """組合完整的 HTML 文件（頁首＋內文＋頁尾），含 CSS 設定。"""
        self.ensure_one()
        body = rendered_body or self.get_content_html()
        header = self.header_html or ''
        footer = self.footer_html or ''

        page_sizes = {
            'A4': ('210mm', '297mm'),
            'A3': ('297mm', '420mm'),
            'A5': ('148mm', '210mm'),
            'letter': ('215.9mm', '279.4mm'),
            'legal': ('215.9mm', '355.6mm'),
        }
        w, h = page_sizes.get(self.page_format, ('210mm', '297mm'))

        return f"""<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<style>
@page {{
    size: {w} {h};
}}
body {{
    font-family: 'Microsoft JhengHei', 'Noto Sans TC', 'Arial', sans-serif;
    font-size: 12pt;
    line-height: 1.6;
    color: #333;
    margin: 0;
    padding: 0;
}}
.doc-header {{ margin-bottom: 12px; border-bottom: 1px solid #ddd; padding-bottom: 8px; }}
.doc-footer {{ margin-top: 12px; border-top: 1px solid #ddd; padding-top: 8px; font-size: 10pt; }}
.doc-page-break {{ page-break-after: always; }}
.doc-field-token {{
    background: #e3f2fd;
    border: 1px solid #90caf9;
    border-radius: 3px;
    padding: 1px 4px;
    font-family: monospace;
    font-size: 0.875em;
    color: #1565c0;
}}
table {{ border-collapse: collapse; width: 100%; }}
td, th {{ border: 1px solid #ccc; padding: 6px; }}
</style>
</head>
<body>
{'<div class="doc-header">' + header + '</div>' if header else ''}
<div class="doc-body">{body}</div>
{'<div class="doc-footer">' + footer + '</div>' if footer else ''}
</body>
</html>"""

    def _generate_pdf(self, record=None):
        """使用 Odoo 內建 wkhtmltopdf 產生 PDF。"""
        self.ensure_one()
        body = self.get_content_html()
        if record:
            body = self._render_template(body, record)

        full_html = self._build_full_html(rendered_body=body)
        Report = self.env['ir.actions.report']
        # px → mm（96dpi：1px = 25.4/96 mm）
        px_to_mm = 25.4 / 96
        pdf_bytes = Report._run_wkhtmltopdf(
            [full_html],          # str，不預先 encode（_run_wkhtmltopdf 內部自行 encode）
            landscape=False,
            specific_paperformat_args={
                # 傳入實際 mm 值，由 wkhtmltopdf CLI 套用正確邊距。
                # @page CSS 只保留 size，避免 CSS margin 與 CLI margin 疊加。
                # left/right 需依賴本模組 report_overrides.py 覆寫的支援。
                'data-report-margin-top':    round(self.margin_top    * px_to_mm, 1),
                'data-report-margin-bottom': round(self.margin_bottom * px_to_mm, 1),
                'data-report-margin-left':   round(self.margin_left   * px_to_mm, 1),
                'data-report-margin-right':  round(self.margin_right  * px_to_mm, 1),
            },
        )
        return pdf_bytes

    def _wrap_html_for_libreoffice(self, body_html):
        """為 LibreOffice 建立最小化的標準 HTML 封裝。

        包含 @page CSS 以確保 DOCX 邊距與螢幕上的 doc-page-sheet 一致。
        注意：此 HTML 僅供 LibreOffice 中介轉換，不會存入 DB 或顯示給使用者。
        """
        # px → cm（96dpi：1px = 2.54/96 cm）
        px_to_cm = 2.54 / 96
        page_sizes_cm = {
            'A4':     (21.0,   29.7),
            'A3':     (29.7,   42.0),
            'A5':     (14.8,   21.0),
            'letter': (21.59,  27.94),
            'Letter': (21.59,  27.94),
            'legal':  (21.59,  35.56),
            'Legal':  (21.59,  35.56),
        }
        w_cm, h_cm = page_sizes_cm.get(self.page_format, (21.0, 29.7))
        mt = round(self.margin_top    * px_to_cm, 2)
        mr = round(self.margin_right  * px_to_cm, 2)
        mb = round(self.margin_bottom * px_to_cm, 2)
        ml = round(self.margin_left   * px_to_cm, 2)
        usable_w_cm = round(w_cm - ml - mr, 2)
        return f"""<html xmlns:o="urn:schemas-microsoft-com:office:office"
  xmlns:w="urn:schemas-microsoft-com:office:word"
  xmlns="http://www.w3.org/TR/REC-html40">
<head>
<meta charset="utf-8">
<style>
@page {{ size: {w_cm}cm {h_cm}cm; margin: {mt}cm {mr}cm {mb}cm {ml}cm; }}
body {{ font-family: sans-serif; max-width: {usable_w_cm}cm; margin: 0 auto; }}
table {{ border-collapse: collapse; }}
td, th {{ border: 1px solid black; padding: 4px; word-break: break-word; }}
img {{ max-width: 100%; height: auto; }}
</style>
</head>
<body>
{body_html}
</body>
</html>"""

    def _generate_docx_via_libreoffice(self, record=None):
        """使用 LibreOffice headless 將 HTML 轉換為 DOCX。
        LibreOffice 不可用時自動 fallback 至 python-docx。
        """
        self.ensure_one()
        import shutil
        if not shutil.which('soffice'):
            # Fallback：使用 python-docx（已 pip install）
            return self._generate_docx_via_python(record)

        body = self.get_content_html()
        if record:
            body = self._render_template(body, record)

        # 使用專為 LibreOffice 設計的最小化 HTML 封裝
        # （不用 _build_full_html，避免 @page CSS 觸發 LO MIME 誤判）
        lo_html = self._wrap_html_for_libreoffice(body)

        with tempfile.TemporaryDirectory() as tmpdir:
            html_path = os.path.join(tmpdir, 'input.html')
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(lo_html)

            result = subprocess.run(
                [
                    'soffice', '--headless', '--norestore',
                    # .html 副檔名已足夠讓 LO 識別輸入格式（不需 --infilter）
                    # 'MS Word 2007 XML' 是 LO 的 DOCX export filter 正式名稱，
                    # 明確指定可防止 LO 找不到 export filter 而失敗
                    '--convert-to', 'docx:MS Word 2007 XML',
                    '--outdir', tmpdir,
                    html_path,
                ],
                capture_output=True,
                timeout=60,
            )
            if result.returncode != 0:
                raise UserError(
                    f'LibreOffice 轉換失敗：{result.stderr.decode("utf-8", errors="replace")}'
                )

            docx_path = os.path.join(tmpdir, 'input.docx')
            if not os.path.exists(docx_path):
                # LO 有時以原始檔名為基礎產生輸出，嘗試搜尋
                candidates = [
                    f for f in os.listdir(tmpdir)
                    if f.endswith('.docx')
                ]
                if candidates:
                    docx_path = os.path.join(tmpdir, candidates[0])
                else:
                    raise UserError(
                        'LibreOffice 轉換後找不到輸出檔案。'
                        f'（stderr: {result.stderr.decode("utf-8", errors="replace")[:500]}）'
                    )

            with open(docx_path, 'rb') as f:
                raw_bytes = f.read()

        # 後處理：修正 LO 以內建預設邊距計算的表格寬度（200mm → 實際版心寬度）
        px_to_mm = 25.4 / 96
        page_w_map = {
            'A4': 210, 'A3': 297, 'A5': 148,
            'letter': 215.9, 'Letter': 215.9,
            'legal': 215.9, 'Legal': 215.9,
        }
        w_mm = page_w_map.get(self.page_format, 210)
        usable_mm = w_mm - (self.margin_left * px_to_mm) - (self.margin_right * px_to_mm)
        return _fix_docx_table_widths(raw_bytes, usable_mm)

    def _generate_docx_via_python(self, record=None):
        """使用 python-docx + lxml 將 HTML 轉換為 DOCX（LibreOffice 不可用時的 fallback）。"""
        self.ensure_one()
        try:
            from docx import Document
            from docx.shared import Mm
            from lxml import html as lhtml
        except ImportError as e:
            raise UserError(
                f'無法匯出 DOCX：{e}\n'
                '請安裝 python-docx：pip install python-docx'
            )

        body_html = self.get_content_html()
        if record:
            body_html = self._render_template(body_html, record)

        doc = Document()

        # ── 設定頁面大小與邊距 ──
        page_sizes_mm = {
            'A4':     (210, 297),
            'A3':     (297, 420),
            'A5':     (148, 210),
            'letter': (216, 279),
            'legal':  (216, 356),
        }
        w_mm, h_mm = page_sizes_mm.get(self.page_format, (210, 297))
        px_to_mm = 0.264583  # 96dpi：1px = 0.264583mm

        section = doc.sections[0]
        section.page_width   = Mm(w_mm)
        section.page_height  = Mm(h_mm)
        section.top_margin   = Mm(self.margin_top    * px_to_mm)
        section.bottom_margin = Mm(self.margin_bottom * px_to_mm)
        section.left_margin  = Mm(self.margin_left   * px_to_mm)
        section.right_margin = Mm(self.margin_right  * px_to_mm)

        # ── 解析 HTML 並轉換結構 ──
        try:
            tree = lhtml.fromstring(f'<div>{body_html}</div>')
            _html_node_to_docx(doc, tree)
        except Exception:
            # 解析失敗時退回純文字
            from lxml import html as lhtml2
            text = lhtml2.fromstring(f'<div>{body_html}</div>').text_content()
            doc.add_paragraph(text)

        # 若頁首/頁尾有內容，加入文件尾部（簡單實作）
        if self.header_html:
            from lxml import html as lhtml3
            header_text = lhtml3.fromstring(f'<div>{self.header_html}</div>').text_content()
            if header_text.strip():
                para = doc.sections[0].header.paragraphs[0]
                para.text = header_text.strip()
        if self.footer_html:
            from lxml import html as lhtml4
            footer_text = lhtml4.fromstring(f'<div>{self.footer_html}</div>').text_content()
            if footer_text.strip():
                para = doc.sections[0].footer.paragraphs[0]
                para.text = footer_text.strip()

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ─── 版本快照 ────────────────────────────────────────────────────

    def action_save_version(self, label=None):
        """在 mail thread 中儲存版本快照。"""
        self.ensure_one()
        html = self.get_content_html()
        self.message_post(
            body=f'<div class="doc-version-snapshot">'
                 f'<strong>版本快照</strong> {label or ""}<br/>{html}</div>',
            subtype_xmlid='mail.mt_note',
            message_type='comment',
        )

    def get_version_list(self):
        """取得版本快照清單。"""
        self.ensure_one()
        messages = self.message_ids.filtered(
            lambda m: '<div class="doc-version-snapshot"' in (m.body or '')
        )
        return [
            {'id': msg.id, 'date': str(msg.date), 'author': msg.author_id.name}
            for msg in messages.sorted('date', reverse=True)
        ]

    # ─── DB 索引 ─────────────────────────────────────────────────────

    def init(self):
        self.env.cr.execute("""
            CREATE INDEX IF NOT EXISTS doc_document_company_write_date_idx
            ON doc_document (company_id, write_date DESC);
            CREATE INDEX IF NOT EXISTS doc_document_company_model_idx
            ON doc_document (company_id, model_id) WHERE model_id IS NOT NULL;
        """)
